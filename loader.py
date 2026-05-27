"""
loader.py — Document loading and chunking
Supports: PDF, TXT, DOCX from uploaded Streamlit files or local paths
"""

import io
import os
from typing import List

from langchain_core.documents import Document

SUPPORTED_EXTENSIONS = (".pdf", ".txt", ".docx")


def load_uploaded_file(uploaded_file) -> str:
    """Load text from a Streamlit UploadedFile object (PDF, TXT, DOCX)."""
    filename = uploaded_file.name.lower()
    content = uploaded_file.getvalue()
    buffer = io.BytesIO(content)
    buffer.name = uploaded_file.name

    if filename.endswith(".txt"):
        return _load_txt(buffer)
    elif filename.endswith(".pdf"):
        return _load_pdf(buffer)
    elif filename.endswith(".docx"):
        return _load_docx(buffer)
    else:
        raise ValueError(f"Unsupported file type: {uploaded_file.name}")


def load_local_file(filepath: str) -> str:
    """Load text from a local file path."""
    ext = os.path.splitext(filepath)[1].lower()
    with open(filepath, "rb") as f:
        content = f.read()
    buf = io.BytesIO(content)
    buf.name = os.path.basename(filepath)

    if ext == ".txt":
        return content.decode("utf-8", errors="ignore")
    elif ext == ".pdf":
        return _load_pdf(buf)
    elif ext == ".docx":
        return _load_docx(buf)
    else:
        raise ValueError(f"Unsupported file type: {filepath}")


def load_local_directory(directory: str) -> List[Document]:
    """Load all supported files from a local directory."""
    docs = []
    if not os.path.isdir(directory):
        raise FileNotFoundError(f"Directory not found: {directory}")

    for fname in os.listdir(directory):
        if fname.lower().endswith(SUPPORTED_EXTENSIONS):
            fpath = os.path.join(directory, fname)
            try:
                text = load_local_file(fpath)
                docs.append(
                    Document(
                        page_content=text,
                        metadata={"source": fname, "path": fpath},
                    )
                )
            except Exception as e:
                print(f"[loader] Skipping {fname}: {e}")
    return docs


def load_uploaded_documents(uploaded_files) -> List[Document]:
    """
    Process uploaded Streamlit files into document objects.
    """
    try:
        import streamlit as st
    except ImportError:
        st = None

    documents = []
    for uf in uploaded_files:
        try:
            text = load_uploaded_file(uf)
            documents.append(
                Document(
                    page_content=text,
                    metadata={"source": uf.name, "origin": "upload"},
                )
            )
        except Exception as e:
            if st is not None:
                st.warning(f"Could not process {uf.name}: {e}")
            else:
                print(f"[loader] Could not process {uf.name}: {e}")
    return documents


# ── Private helpers ──────────────────────────────────────────────────────────

def _load_txt(file_obj) -> str:
    raw = file_obj.read()
    if isinstance(raw, bytes):
        return raw.decode("utf-8", errors="ignore")
    return raw


def _load_pdf(file_obj) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("Install pypdf:  pip install pypdf")

    reader = PdfReader(file_obj)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _load_docx(file_obj) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Install python-docx:  pip install python-docx")

    doc = Document(file_obj)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)

def load_documents(uploaded_files=None, directory="data"):
    """
    Universal loader:
    - If uploaded_files → use Streamlit uploads
    - Else → load from local directory
    """

    # Case 1: UI upload
    if uploaded_files:
        return load_uploaded_documents(uploaded_files)

    # Case 2: CLI / local files
    return load_local_directory(directory)
